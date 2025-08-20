import streamlit as st
from docx import Document
from docx.shared import Pt
import pandas as pd
import numpy as np
import os
# Streamlit UI
st.title("Fill Word Document Form")

with st.form("user_form"):
    name = st.text_input("Full Name")
    your_gender = st.selectbox("Gender", ["Male", "Female", "Other"])
    phone1 = st.number_input("Phone Number", min_value=1000000000, max_value=9999999999,step=1, format="%d")
    phone2 = st.number_input("Alternate Phone Number", min_value=1000000000, max_value=9999999999, step=1, format="%d")
    Address=st.text_input("Address")
    email1 = st.text_input("Email")
    pincode = st.text_input("Pincode")
    district = st.text_input("District")
    state = st.text_input("State")

    submitted = st.form_submit_button("Generate Document")

if submitted:
    doc = Document("tf.docx")

    # full name
    full_name = doc.paragraphs[5]
    full_name_text = full_name.text
    # print(full_name_text)
    full_name.clear()  # keeping as-is
    full_name_run = full_name.add_run("Full Name:")
    full_name_answer_run = full_name.add_run(" <<full name>>")
    full_name_answer_run.font.size = Pt(12)
    full_name_answer_run.font.bold = True
    full_name_answer_run.text = full_name_answer_run.text.replace("<<full name>>", name)

    # gender
    gender = doc.paragraphs[6]
    gender_text = doc.paragraphs[6].text.split()
    gender_text1, gender_text2 = gender_text[0], gender_text[1]
    gender.clear()
    gender_field_run = gender.add_run("Gender:")
    gender_answer_run = gender.add_run(" " + "<<gender>>")
    gender_answer_run.font.bold = True
    gender_answer_run.font.size = Pt(12)
    gender_answer_run.text = gender_answer_run.text.replace("<<gender>>", your_gender)

    # pincode, district, state
    pincode_district_state = doc.paragraphs[11]
    pincode_district_state_text = doc.paragraphs[11].text.split()
    pincode_district_state.clear()
    pincode_field_run = pincode_district_state.add_run('Pincode:')
    pincode_field_answer_run = pincode_district_state.add_run(' ' + '<<pincode>>')
    pincode_field_answer_run.font.size = Pt(12)
    pincode_field_answer_run.font.bold = True
    pincode_field_answer_run.text = pincode_field_answer_run.text.replace('<<pincode>>', pincode)

    district_field_run = pincode_district_state.add_run('             District:')
    district_field_answer_run = pincode_district_state.add_run(' ' + '<<district>>')
    district_field_answer_run.font.size = Pt(12)
    district_field_answer_run.font.bold = True
    district_field_answer_run.text = district_field_answer_run.text.replace('<<district>>', district)

    state_field_run = pincode_district_state.add_run('       State:')
    state_field_answer_run = pincode_district_state.add_run(' ' + '<<state>>')
    state_field_answer_run.font.size = Pt(12)
    state_field_answer_run.font.bold = True
    state_field_answer_run.text = state_field_answer_run.text.replace('<<state>>', state)

    # phone numbers
    phone_1_2 = doc.paragraphs[7]
    phone_1_2_text = doc.paragraphs[7].text.split()
    phone_1_2.clear()
    phone_1_field_run = phone_1_2.add_run('Phone Number:')
    phone_1_field_answer_run = phone_1_2.add_run(' ' + '<<phone1>>')
    phone_1_field_answer_run.font.bold = True
    phone_1_field_answer_run.font.size = Pt(12)
    phone_1_field_answer_run.text = phone_1_field_answer_run.text.replace("<<phone1>>", str(phone1))

    phone_2_field_run = phone_1_2.add_run('          Alternate Number:')
    phone_2_field_answer_run = phone_1_2.add_run(' ' + '<<phone2>>')
    phone_2_field_answer_run.font.bold = True
    phone_2_field_answer_run.font.size = Pt(12)
    phone_2_field_answer_run.text = phone_2_field_answer_run.text.replace("<<phone2>>", str(phone2))

    # email
    email = doc.paragraphs[8]
    email_text = doc.paragraphs[8].text
    email.clear()
    email_field_run = email.add_run("Email:")
    email_field_answer_run = email.add_run(" " + "<<email>>")
    email_field_answer_run.font.bold = True
    email_field_answer_run.font.size = Pt(12)
    email_field_answer_run.text = email_field_answer_run.text.replace("<<email>>", email1)

    # Save updated document
    output_file = f"{name}.docx"
    doc.save(output_file)

    # ✅ Save to Excel
    new_data = {
        "Name": name,
        "Gender": your_gender,
        "Phone Number": phone1,
        "Alternate Phone Number": phone2,
        "Email": email1,
        "Pincode": pincode,
        "District": district,
        "State": state,
        "Address":Address
    }

    excel_file = "data.xlsx"
        # ✅ Append to Excel safely
    if os.path.exists(excel_file):
        df = pd.read_excel(excel_file)
        df = pd.concat([df, pd.DataFrame([new_data])], ignore_index=True)
    else:
        df = pd.DataFrame([new_data])

    df.to_excel(excel_file, index=False)

    # ✅ Success messages
    st.success(f"✅ Document generated: {output_file}")
    st.success(f"✅ Data saved to Excel: {excel_file}")

    # ✅ Show updated table
    # st.subheader("📊 Updated Excel Data")
    